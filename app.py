import json
import os
import re
import sqlite3
import hashlib
from datetime import datetime
from pathlib import Path

import pandas as pd
import streamlit as st
from PyPDF2 import PdfReader
import docx

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from pdf2image import convert_from_path
    import pytesseract
except ImportError:
    convert_from_path = None
    pytesseract = None


def sanitize_filename(filename):
    cleaned = re.sub(r"[^A-Za-z0-9._-]+", "_", filename)
    return cleaned.strip("_.")


def highlight_term(raw_text, term):
    if not term or not raw_text:
        return raw_text
    term_list = [w for w in str(term).split() if w.strip()]
    if not term_list:
        return raw_text
    safe_terms = [re.escape(w) for w in term_list]
    pattern = re.compile(rf"({'|'.join(safe_terms)})", re.IGNORECASE)
    highlighted = pattern.sub(r"<mark>\1</mark>", raw_text)
    return highlighted

DB_PATH = Path("filedata.db")
UPLOAD_DIR = Path("uploads")
UPLOAD_DIR.mkdir(exist_ok=True)

def init_db():
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute(
        """
        CREATE TABLE IF NOT EXISTS files (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            filename TEXT,
            stored_path TEXT,
            mime_type TEXT,
            size INTEGER,
            uploaded_at TEXT,
            sha256 TEXT UNIQUE,
            metadata TEXT,
            full_text TEXT
        )
        """
    )
    conn.commit()
    conn.close()


def compute_sha256(file_path):
    h = hashlib.sha256()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return h.hexdigest()


def extract_text(file_path):
    ext = file_path.suffix.lower()
    if ext == ".txt":
        return file_path.read_text(errors="ignore")
    if ext == ".pdf":
        text = ""
        try:
            reader = PdfReader(str(file_path))
            pages = []
            for page in reader.pages:
                pages.append(page.extract_text() or "")
            text = "\n".join(pages)
        except Exception:
            text = ""

        if not text and pdfplumber is not None:
            try:
                with pdfplumber.open(str(file_path)) as pdf:
                    pages = [p.extract_text() or "" for p in pdf.pages]
                text = "\n".join(pages)
            except Exception:
                text = ""

        if not text and convert_from_path is not None and pytesseract is not None:
            try:
                images = convert_from_path(str(file_path), dpi=200)
                ocr_text = []
                for img in images:
                    ocr_text.append(pytesseract.image_to_string(img))
                text = "\n".join(ocr_text)
            except Exception:
                text = ""

        return text
    if ext == ".docx":
        try:
            doc = docx.Document(file_path)
            return "\n".join(p.text for p in doc.paragraphs)
        except Exception:
            return ""
    return ""


def extract_metadata(file_path):
    data = {}
    stt = file_path.stat()
    data.update(
        {
            "size": stt.st_size,
            "created_at": datetime.fromtimestamp(stt.st_ctime).isoformat(),
            "modified_at": datetime.fromtimestamp(stt.st_mtime).isoformat(),
        }
    )
    ext = file_path.suffix.lower()
    if ext == ".pdf":
        try:
            reader = PdfReader(str(file_path))
            if reader.metadata:
                data["pdf_metadata"] = {k: v for k, v in reader.metadata.items()}
        except Exception:
            pass
    if ext == ".docx":
        try:
            doc = docx.Document(file_path)
            props = doc.core_properties
            data["docx_metadata"] = {
                "author": props.author,
                "title": props.title,
                "subject": props.subject,
                "last_modified_by": props.last_modified_by,
                "created": props.created.isoformat() if props.created else None,
                "modified": props.modified.isoformat() if props.modified else None,
            }
        except Exception:
            pass
    return data


def insert_file_record(filename, stored_path, mime_type, size, sha256, metadata, full_text):
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute(
        """
        INSERT OR IGNORE INTO files
        (filename, stored_path, mime_type, size, uploaded_at, sha256, metadata, full_text)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            filename,
            str(stored_path),
            mime_type,
            size,
            datetime.utcnow().isoformat(),
            sha256,
            json.dumps(metadata, ensure_ascii=False),
            full_text,
        ),
    )
    conn.commit()
    conn.close()


def main():
    st.title("Metadata and Text Extractor")
    init_db()

    tab1, tab2 = st.tabs(["Upload file", "Search files"])

    with tab1:
        st.header("Upload and index a file")
        file_uploader = st.file_uploader("Select file", type=None)

        if file_uploader is not None:
            if st.button("Upload and index"):
                safe_name = sanitize_filename(file_uploader.name)
                if not safe_name:
                    safe_name = "uploaded_file"
                target = UPLOAD_DIR / f"{datetime.utcnow().strftime('%Y%m%d%H%M%S')}_{safe_name}"
                with open(target, "wb") as out:
                    out.write(file_uploader.getbuffer())
                sha256 = compute_sha256(target)
                text = extract_text(target)
                metadata = extract_metadata(target)
                insert_file_record(
                    file_uploader.name,
                    target,
                    file_uploader.type or "application/octet-stream",
                    target.stat().st_size,
                    sha256,
                    metadata,
                    text,
                )
                st.success("File uploaded and indexed successfully.")
                st.write("SHA256:", sha256)
                if text:
                    st.subheader("Extracted text")
                    st.text_area("File text", value=text, height=700)
                else:
                    st.warning("Text extraction did not return content. If this is a scanned PDF, consider adding OCR support in requirements.")


    with tab2:
        st.header("Search uploaded files")

        if st.button("Cleanup missing local files from index"):
            conn = sqlite3.connect(DB_PATH)
            c = conn.cursor()
            c.execute("SELECT id, stored_path FROM files")
            orphan_ids = []
            for _id, stored_path in c.fetchall():
                if not Path(stored_path).exists():
                    orphan_ids.append(_id)
            if orphan_ids:
                c.execute(f"DELETE FROM files WHERE id IN ({','.join(['?']*len(orphan_ids))})", orphan_ids)
                conn.commit()
            conn.close()
            st.success(f"Removed {len(orphan_ids)} orphaned index entries.")

        view_mode = st.selectbox("Result view", ["Grid", "List"])
        q = st.text_input("Search query (name, metadata, text)")
        if q:
            conn = sqlite3.connect(DB_PATH)
            c = conn.cursor()
            term = f"%{q}%"
            c.execute(
                """
                SELECT id, filename, stored_path, mime_type, size, uploaded_at, metadata, full_text
                FROM files
                WHERE filename LIKE ? OR metadata LIKE ? OR full_text LIKE ?
                ORDER BY uploaded_at DESC
                """,
                (term, term, term),
            )
            rows = c.fetchall()
            conn.close()

            st.write(f"{len(rows)} files found")
            if rows:
                if view_mode == "List":
                    # List mode uses cards
                    cols_count = 3
                    rows_by_group = [rows[i : i + cols_count] for i in range(0, len(rows), cols_count)]
                    for group in rows_by_group:
                        cols = st.columns(len(group))
                        for col, row in zip(cols, group):
                            _id, filename, stored_path, mime_type, size, uploaded_at, metadata_raw, full_text = row
                            meta = json.loads(metadata_raw or "{}")
                            col.markdown(f"**{filename}**")
                            col.markdown(f"Type: `{mime_type}`")
                            col.markdown(f"Size: `{size}` bytes")
                            col.markdown(f"Uploaded: `{uploaded_at}`")
                            col.markdown(f"Metadata: `{json.dumps(meta, ensure_ascii=False)[:120]}...`")
                            full_highlighted = highlight_term((full_text or "").replace("\n", " ").strip(), q)
                            if full_highlighted:
                                with col.expander("Show extracted text"):
                                    col.markdown(full_highlighted, unsafe_allow_html=True)
                            else:
                                col.markdown("*No extracted text available*. Use an exact keyword search to verify content, or re-upload the file with a clear text PDF or use OCR.")

                            stored_path_obj = Path(stored_path)
                            if stored_path_obj.exists():
                                col.download_button(
                                    label="Download",
                                    data=open(stored_path_obj, "rb").read(),
                                    file_name=filename,
                                    mime=mime_type,
                                )
                            else:
                                col.warning("File not found on disk (it may have been moved or deleted).")
                else:
                    # Grid mode uses table
                    results = []
                    for row in rows:
                        _id, filename, stored_path, mime_type, size, uploaded_at, metadata_raw, full_text = row
                        meta = json.loads(metadata_raw or "{}")
                        snippet = (full_text or "").replace("\n", " ").strip()[:180]
                        highlighted = highlight_term(snippet, q)
                        results.append({
                            "Filename": filename,
                            "Type": mime_type,
                            "Size(Bytes)": size,
                            "Uploaded At": uploaded_at,
                            "Metadata": json.dumps(meta, ensure_ascii=False)[:100],
                        })

                    # Table with highlights via HTML
                    html_table = pd.DataFrame(results).to_html(index=False, escape=False)
                    st.markdown(html_table, unsafe_allow_html=True)

                    # show extracted text and action links per found row
                    for row in rows:
                        _, filename, stored_path, mime_type, _, _, _, full_text = row
                        if full_text:
                            highlighted_full = highlight_term(full_text, q)
                            with st.expander(f"Extracted text for {filename}"):
                                st.markdown(highlighted_full, unsafe_allow_html=True)
                        else:
                            st.write(f"No extracted text available for {filename}.")

                        stored_path_obj = Path(stored_path)
                        if stored_path_obj.exists():
                            st.download_button(
                                label=f"Download {filename}",
                                data=open(stored_path_obj, "rb").read(),
                                file_name=filename,
                                mime=mime_type,
                            )
                        else:
                            st.warning(f"File not found on disk for {filename} (may have been moved or deleted).")
        else:
            st.info("Enter search text to filter uploaded files.")


if __name__ == "__main__":
    main()
